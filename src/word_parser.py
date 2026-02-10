import logging
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from src.file_parser import FileParser

logger = logging.getLogger(__name__)


class WordParser(FileParser, extensions=(".docx",)):
    """Extracts text from Word (.docx) resume files."""

    def parse(self, file_path: Path) -> str:
        """Extract text from a .docx file.

        Raises:
            RuntimeError: If the file is corrupt, not a valid .docx, or
                cannot be read.
        """
        try:
            doc = Document(str(file_path))
        except PackageNotFoundError as exc:
            logger.error("Not a valid .docx package: '%s'", file_path.name)
            raise RuntimeError(
                f"'{file_path.name}' is not a valid Word document. "
                f"The file may be corrupt or in an older .doc format."
            ) from exc
        except zipfile.BadZipFile as exc:
            logger.error("Bad zip structure in '%s': %s", file_path.name, exc)
            raise RuntimeError(
                f"'{file_path.name}' has an invalid archive structure. "
                f"The file may be corrupt or not a real .docx file."
            ) from exc
        except PermissionError as exc:
            logger.error("Permission denied reading '%s': %s", file_path.name, exc)
            raise
        except Exception as exc:
            logger.error("Failed to open Word file '%s': %s", file_path.name, exc)
            raise RuntimeError(f"Cannot open '{file_path.name}': {exc}") from exc

        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            parts.append(text)

        if parts:
            logger.debug("Extracted %d text parts from paragraphs/tables", len(parts))
            return "\n".join(parts)

        logger.info("Paragraphs/tables empty — falling back to raw XML extraction")
        return self._extract_from_xml(file_path)

    @staticmethod
    def _extract_from_xml(file_path: Path) -> str:
        """Extract text directly from the document.xml inside the .docx zip."""
        try:
            with zipfile.ZipFile(file_path) as zf:
                if "word/document.xml" not in zf.namelist():
                    logger.error("word/document.xml not found in '%s'", file_path.name)
                    raise RuntimeError(
                        f"'{file_path.name}' does not contain a valid Word "
                        f"document structure (missing word/document.xml)."
                    )
                xml_bytes = zf.read("word/document.xml")
        except zipfile.BadZipFile as exc:
            logger.error("Bad zip in XML fallback for '%s': %s", file_path.name, exc)
            raise RuntimeError(
                f"'{file_path.name}' has an invalid archive structure."
            ) from exc

        xml_text = xml_bytes.decode("utf-8")
        fragments = re.findall(r"<w:t[^>]*>([^<]+)</w:t>", xml_text)
        return " ".join(fragments)
